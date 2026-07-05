import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re

# إعدادات واجهة المستخدم
st.set_page_config(page_title="نظام تنسيق وتدقيق كشوفات الوكلاء", layout="wide")
st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { background-color: #2E4053; color: white; width: 100%; font-weight: bold; border-radius: 8px; font-size: 18px;}
    .report-box { background-color: #F4F6F7; padding: 15px; border-radius: 8px; border-right: 5px solid #2E4053; text-align: right; margin-bottom: 10px;}
    div.row-widget.stRadio > div { flex-direction: row-reverse; justify-content: flex-start; gap: 20px; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: right;'>نظام تنسيق وتدقيق كشوفات الوكلاء المطور 📄💎</h1>", unsafe_allow_html=True)

if "processing_done" not in st.session_state:
    st.session_state.processing_done = False
    st.session_state.df_final = None
    st.session_state.output_filename = ""
    st.session_state.selected_card = ""
    st.session_state.template_choice = ""
    st.session_state.name_format = "" 
    st.session_state.uploaded_filenames = []

# -----------------------------------------------------------------------------
# مساعدات التنسيق المتقدمة لملفات Word عبر الـ XML
# -----------------------------------------------------------------------------
def set_table_borders(table, color_hex="2A4B7C"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_vertical_text(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    text_dir = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="btLr"/>')
    tcPr.append(text_dir)

def set_cell_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = parse_xml(f'<w:noWrap {nsdecls("w")}/>')
    tcPr.append(no_wrap)

def format_cell_advanced(cell, text, bold=False, color_rgb=None, size_pt=16, font_name="Calibri", align="center"):
    cell.text = str(text)
    p = cell.paragraphs[0]
    
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    
    for run in p.runs:
        run.bold = bold
        if color_rgb:
            run.font.color.rgb = color_rgb
            
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        run.font.size = Pt(size_pt)

# -----------------------------------------------------------------------------
# محرك قراءة وتنظيف البيانات المطور (يدعم الرفع المتعدد والدمج)
# -----------------------------------------------------------------------------
def extract_and_clean_data(file_objs, card_choice, name_format):
    raw_records = []
    
    # حلقة تكرارية لمعالجة كل ملف تم رفعه
    for file_obj in file_objs:
        rows_data = []
        file_ext = file_obj.name.split('.')[-1].lower()
        
        if file_ext == 'docx':
            doc = Document(file_obj)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    rows_data.append(cells)
                    
        elif file_ext == 'xlsx':
            xls = pd.ExcelFile(file_obj)
            for sheet_name in xls.sheet_names:
                df_excel = pd.read_excel(xls, sheet_name=sheet_name, header=None)
                for row in df_excel.values:
                    cells = []
                    for cell in row:
                        if pd.isna(cell):
                            continue
                        if isinstance(cell, float) and cell.is_integer():
                            cells.append(str(int(cell)))
                        else:
                            cells.append(str(cell).strip().replace('\n', ' '))
                    rows_data.append(cells)
        
        # استخراج البيانات من السطور
        for cells in rows_data:
            if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells):
                continue
            
            name_idx = -1
            max_len = 0
            for i, c in enumerate(cells):
                if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                    if len(c) > max_len:
                        max_len = len(c)
                        name_idx = i
            if name_idx == -1: continue
            
            card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
            if not card_indices: continue
            
            old_card_num = cells[card_indices[0]]
            new_card_num = cells[card_indices[1]] if len(card_indices) > 1 else old_card_num
            selected_card_num = new_card_num if card_choice == "رقم البطاقة الحديث" else old_card_num
            
            digit_cells = [int(cells[i]) for i in range(name_idx) if cells[i].isdigit()]
            if len(digit_cells) >= 3:
                withheld, eligible, total = digit_cells[0], digit_cells[1], digit_cells[2]
            elif len(digit_cells) == 2:
                withheld, eligible, total = 0, digit_cells[0], digit_cells[1]
            else:
                continue
            
            full_name = cells[name_idx].strip()
            if name_format == "الاسم الثلاثي فقط (بدون لقب)":
                name_parts = full_name.split()
                final_name = " ".join(name_parts[:3])
            else:
                final_name = full_name 
                
            raw_records.append({
                "اسم رب الأسرة": final_name,
                "رقم البطاقة": selected_card_num,
                "الكلي": total,
                "محجوب": withheld,
                "مستحق": eligible
            })
            
    # الدمج والفرز بعد الانتهاء من جميع الملفات
    df = pd.DataFrame(raw_records)
    if not df.empty:
        # إزالة القيود المكررة تماماً (إن وجدت في أكثر من ملف)
        df = df.drop_duplicates(subset=["اسم رب الأسرة", "رقم البطاقة", "الكلي"])
        # الفرز الأبجدي بناءً على اسم رب الأسرة
        df = df.sort_values(by="اسم رب الأسرة").reset_index(drop=True)
        # إعادة إضافة التسلسل الرقمي الموحد (ت)
        df.insert(0, "ت", df.index + 1)
        
    return df

# -----------------------------------------------------------------------------
# محرك بناء تقرير Word - النموذج الأول (الأصلي)
# -----------------------------------------------------------------------------
def build_professional_word_report(df, filename_base, card_choice):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.3)
        section.right_margin = Cm(0.3)
        
    clean_name = filename_base
    for w in ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]:
        clean_name = clean_name.replace(w, "")
    clean_name = re.sub(r'[a-zA-Z]', '', clean_name)
    clean_name = re.sub(r'[\-_+_.]', '', clean_name)
    clean_name = " ".join(clean_name.split())
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_p.add_run(f"الكشف الإحصائي المنسق للوكيل: {clean_name}")
    title_run.font.name = "Segoe UI Semibold"
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    headers = ["ت", "اسم رب الأسرة", "حقل فارغ", "الكلي", "مستحق", "محجوب", card_choice, "ملاحظات"]
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color_hex="2A4B7C")
    table._tbl.tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    table.rows[0]._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    max_name_len = max(df["اسم رب الأسرة"].astype(str).str.len().max(), 15)
    dynamic_name_width = Cm(max_name_len * 0.22 + 0.5)
    col_widths = [Cm(0.9), dynamic_name_width, Cm(0.44), Cm(0.9), Cm(0.9), Cm(0.9), Cm(3.0), Cm(2.19)]
    
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    
    for i, title in enumerate(headers):
        table.rows[0].cells[i].width = col_widths[i]
        if i in [3, 4, 5]:
            set_cell_vertical_text(table.rows[0].cells[i])
            format_cell_advanced(table.rows[0].cells[i], title, bold=True, size_pt=12, font_name="Segoe UI Semibold", align="center", color_rgb=COLOR_NAVY_BLUE)
        else:
            format_cell_advanced(table.rows[0].cells[i], title, bold=True, size_pt=14, font_name="Segoe UI Semibold", align="left" if i==1 else "center", color_rgb=COLOR_NAVY_BLUE)
            
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        table.rows[idx+1]._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        is_eligible_zero = int(row["مستحق"]) == 0
        set_cell_no_wrap(row_cells[1])
        
        for i in range(8):
            row_cells[i].width = col_widths[i]
            val = row["ت"] if i==0 else row["اسم رب الأسرة"] if i==1 else "x" if i==2 and is_eligible_zero else "" if i==2 else row["الكلي"] if i==3 else row["مستحق"] if i==4 else row["محجوب"] if i==5 else row["رقم البطاقة"] if i==6 else "محجوب" if i==7 and is_eligible_zero else ""
            font_size = 14 if i==5 else 12 if i==7 and is_eligible_zero else 16
            text_color = RGBColor(203, 67, 53) if i==7 and is_eligible_zero else None
            format_cell_advanced(row_cells[i], val, size_pt=font_size, font_name="Calibri", color_rgb=text_color, align="left" if i==1 else "center")
            
            if is_eligible_zero: set_cell_background(row_cells[i], "EC7063")
            else:
                if i==0: set_cell_background(row_cells[i], "D4E6F1")
                elif i==3: set_cell_background(row_cells[i], "EBF5FB")
                elif i==4: set_cell_background(row_cells[i], "E8F8F5")
                elif i==5: set_cell_background(row_cells[i], "FADBD8")
    return save_doc_buffer(doc, df)

# -----------------------------------------------------------------------------
# محرك بناء تقرير Word - النموذج الثاني
# -----------------------------------------------------------------------------
def build_professional_word_report_v2(df, filename_base, card_choice):
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Cm(0.5), Cm(0.5), Cm(0.3), Cm(0.3)
        
    clean_name = filename_base
    for w in ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]: clean_name = clean_name.replace(w, "")
    clean_name = " ".join(re.sub(r'[a-zA-Z\-_+_.]', '', clean_name).split())
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_p.add_run(f"الكشف الإحصائي المنسق للوكيل: {clean_name}")
    title_run.font.name, title_run.font.size, title_run.bold = "Segoe UI Semibold", Pt(14), True
    
    headers = ["ت", "اسم رب الأسرة", "حقل فارغ 1", "حقل فارغ 2", "الكلي", "مستحق", "محجوب", card_choice, "ملاحظات"]
    table = doc.add_table(rows=1, cols=9)
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "2A4B7C")
    table._tbl.tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    table.rows[0]._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    dynamic_name_width = Cm(max(df["اسم رب الأسرة"].astype(str).str.len().max(), 15) * 0.22 + 0.5)
    col_widths = [Cm(0.9), dynamic_name_width, Cm(0.80), Cm(0.80), Cm(0.9), Cm(0.9), Cm(0.9), Cm(3.0), Cm(1.80)]
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    
    for i, title in enumerate(headers):
        table.rows[0].cells[i].width = col_widths[i]
        if i in [4, 5, 6]: set_cell_vertical_text(table.rows[0].cells[i])
        format_cell_advanced(table.rows[0].cells[i], title, bold=True, size_pt=14, font_name="Segoe UI Semibold", align="left" if i==1 else "center", color_rgb=COLOR_NAVY_BLUE)
            
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        table.rows[idx+1]._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        is_eligible_zero = int(row["مستحق"]) == 0
        set_cell_no_wrap(row_cells[1])
        
        for i in range(9):
            row_cells[i].width = col_widths[i]
            val = row["ت"] if i==0 else row["اسم رب الأسرة"] if i==1 else "x" if i in [2,3] and is_eligible_zero else "" if i in [2,3] else row["الكلي"] if i==4 else row["مستحق"] if i==5 else row["محجوب"] if i==6 else row["رقم البطاقة"] if i==7 else "محجوب" if i==8 and is_eligible_zero else ""
            text_color = RGBColor(203, 67, 53) if i==8 and is_eligible_zero else None
            format_cell_advanced(row_cells[i], val, size_pt=14, font_name="Calibri", color_rgb=text_color, align="left" if i==1 else "center")
            
            if is_eligible_zero: set_cell_background(row_cells[i], "EC7063")
            else:
                if i==0: set_cell_background(row_cells[i], "D4E6F1")
                elif i==4: set_cell_background(row_cells[i], "EBF5FB")
                elif i==5: set_cell_background(row_cells[i], "E8F8F5")
            if i==6: set_cell_background(row_cells[i], "E5E7E9")
    return save_doc_buffer(doc, df)

# -----------------------------------------------------------------------------
# محرك بناء تقرير Word - النموذج الثالث
# -----------------------------------------------------------------------------
def build_professional_word_report_v3(df, filename_base, card_choice):
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Cm(0.5), Cm(0.5), Cm(0.3), Cm(0.3)
        
    clean_name = filename_base
    for w in ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]:
        clean_name = clean_name.replace(w, "")
    clean_name = re.sub(r'[a-zA-Z]', '', clean_name)
    clean_name = re.sub(r'[\-_+_.]', '', clean_name)
    clean_name = " ".join(clean_name.split())
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_p.add_run(f"الكشف الإحصائي المنسق للوكيل: {clean_name}")
    title_run.font.name, title_run.font.size, title_run.bold = "Segoe UI Semibold", Pt(14), True
    
    headers = ["ت", "اسم رب الأسرة", card_choice, "الكلي", "مستحق", "محجوب", "الشهر الأول", "الشهر الثاني", "الشهر الثالث", "الشهر الرابع"]
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color_hex="2A4B7C")
    table._tbl.tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    table.rows[0]._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    dynamic_name_width = Cm(max(df["اسم رب الأسرة"].astype(str).str.len().max(), 15) * 0.22 + 0.5)
    col_widths = [Cm(0.9), dynamic_name_width, Cm(3.0), Cm(0.9), Cm(0.9), Cm(0.9), Cm(2.3), Cm(2.3), Cm(2.3), Cm(2.3)]
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    
    for i, title in enumerate(headers):
        table.rows[0].cells[i].width = col_widths[i]
        if i in [3, 4, 5]: set_cell_vertical_text(table.rows[0].cells[i])
        format_cell_advanced(table.rows[0].cells[i], title, bold=True, size_pt=12, font_name="Segoe UI Semibold", align="left" if i == 1 else "center", color_rgb=COLOR_NAVY_BLUE)
            
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        table.rows[idx+1]._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        is_eligible_zero = int(row["مستحق"]) == 0
        set_cell_no_wrap(row_cells[1])
        
        for i in range(10):
            row_cells[i].width = col_widths[i]
            val = row["ت"] if i==0 else row["اسم رب الأسرة"] if i==1 else row["رقم البطاقة"] if i==2 else row["الكلي"] if i==3 else row["مستحق"] if i==4 else row["محجوب"] if i==5 else ""
            format_cell_advanced(row_cells[i], val, size_pt=16, font_name="Calibri", color_rgb=None, align="left" if i==1 else "center")
            
            if is_eligible_zero: set_cell_background(row_cells[i], "EC7063")
            else:
                if i == 0: set_cell_background(row_cells[i], "D4E6F1")
            
            if i == 3: set_cell_background(row_cells[i], "E5E7E9")
            if i == 5: set_cell_background(row_cells[i], "FCF3CF")
    return save_doc_buffer(doc, df)

# -----------------------------------------------------------------------------
# محرك بناء تقرير Word - النموذج الرابع
# -----------------------------------------------------------------------------
def build_professional_word_report_v4(df, filename_base, card_choice):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.3)
        section.right_margin = Cm(0.3)
        
    clean_name = filename_base
    words_to_remove = ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]
    for w in words_to_remove:
        clean_name = clean_name.replace(w, "")
    clean_name = re.sub(r'[a-zA-Z]', '', clean_name)
    clean_name = re.sub(r'[\-_+_.]', '', clean_name)
    clean_name = " ".join(clean_name.split())
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_p.add_run(f"الكشف الإحصائي المنسق للوكيل: {clean_name}")
    title_run.font.name = "Segoe UI Semibold"
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    headers = ["ت", card_choice, "اسم المواطن", "العدد الكلي"] + [f"سلة {i}" for i in range(1, 13)]
    
    table = doc.add_table(rows=1, cols=16)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    set_table_borders(table, color_hex="2A4B7C")
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    max_name_len = max(df["اسم رب الأسرة"].astype(str).str.len().max(), 15)
    dynamic_name_width = Cm(max_name_len * 0.22 + 0.5)
    
    col_widths = [Cm(0.9), Cm(2.5), dynamic_name_width, Cm(0.9)] + [Cm(1.05)] * 12 
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        if i >= 3:
            set_cell_vertical_text(hdr_cells[i])
        cell_align = "left" if i == 2 else "center"
        format_cell_advanced(hdr_cells[i], title, bold=True, size_pt=12, font_name="Segoe UI Semibold", align=cell_align, color_rgb=COLOR_NAVY_BLUE)
            
    HEX_ELEGANT_BLUE = "D4E6F1"
    HEX_LIGHT_GREEN = "E8F8F5"
    HEX_ALERT_RED = "EC7063"
    
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        r_trPr = table.rows[idx+1]._tr.get_or_add_trPr()
        r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        is_eligible_zero = int(row["مستحق"]) == 0
        set_cell_no_wrap(row_cells[2])
        
        for i in range(16):
            row_cells[i].width = col_widths[i]
            val = ""
            cell_align = "center"
            font_size = 14
            
            if i == 0: val = row["ت"]
            elif i == 1: val = row["رقم البطاقة"]
            elif i == 2: 
                val = row["اسم رب الأسرة"]
                cell_align = "left" 
            elif i == 3: val = row["الكلي"]
            elif i >= 4: val = ""
                    
            format_cell_advanced(row_cells[i], val, size_pt=font_size, font_name="Calibri", color_rgb=None, align=cell_align)
            
            if is_eligible_zero:
                set_cell_background(row_cells[i], HEX_ALERT_RED)
            else:
                if i == 0: set_cell_background(row_cells[i], HEX_ELEGANT_BLUE)
                if i == 3: set_cell_background(row_cells[i], HEX_LIGHT_GREEN)

    return save_doc_buffer(doc, df)

def save_doc_buffer(doc, df):
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    total_all = df["الكلي"].astype(int).sum()
    total_eligible = df["مستحق"].astype(int).sum()
    total_withheld = df["محجوب"].astype(int).sum()
    
    doc.add_paragraph()  
    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stats_p.paragraph_format.element.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    
    stats_text = (
        f"العدد الكلي للافراد = {total_all}\n"
        f"العدد الكلي للمستحقين = {total_eligible}\n"
        f"العدد الكلي للمحجوبين = {total_withheld}"
    )
    stats_text_run = stats_p.add_run(stats_text)
    stats_text_run.font.name = "Segoe UI Semibold"
    stats_text_run.font.size = Pt(13)
    stats_text_run.bold = True
    stats_text_run.font.color.rgb = COLOR_NAVY_BLUE

    footer = doc.sections[0].footer
    if len(footer.paragraphs) == 0:
        footer_p = footer.add_paragraph()
    else:
        footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.clear()
    f_run = footer_p.add_run("صفحة ")
    f_run.font.size = Pt(10)
    f_run._r.extend([
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'),
        parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'),
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'),
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    ])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# واجهة استخدام التطبيق (تم تفعيل الرفع المتعدد هنا)
# -----------------------------------------------------------------------------
st.markdown("<h3 style='text-align: right;'>📂 رفع كشوفات الوكلاء (يمكنك اختيار أكثر من ملف معاً للدمج)</h3>", unsafe_allow_html=True)
# إضافة accept_multiple_files=True للرفع المتعدد
uploaded_files = st.file_uploader("ارفع كشف أو عدة كشوفات (Word/Excel)", type=['docx', 'xlsx'], accept_multiple_files=True, key="doc_input_v7", label_visibility="collapsed")

st.markdown("<br>", unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)

with col1:
    selected_card = st.radio(
        "📄 اختر نوع رقم البطاقة:",
        ["رقم البطاقة القديم", "رقم البطاقة الحديث"],
        index=0,
        horizontal=False
    )

with col2:
    name_format = st.radio(
        "📝 تنسيق الاسم المستخرج:",
        ["الاسم الثلاثي فقط (بدون لقب)", "الاسم الكامل (مع اللقب)"],
        index=0,
        horizontal=False
    )

with col3:
    template_choice = st.radio(
        "🎨 اختر نموذج قالب الـ Word المطلوب:",
        [
            "النموذج الأول (الأصلي المطور)", 
            "النموذج الثاني (حجم 14 وحقلين فارغين)",
            "النموذج الثالث (خط 16، عناوين 12، 4 أشهر)",
            "النموذج الرابع (12 سلة، العدد الكلي)"
        ],
        index=0,
        horizontal=False
    )

st.markdown("<br>", unsafe_allow_html=True)

if uploaded_files:
    current_filenames = [f.name for f in uploaded_files]
    if (st.session_state.get('uploaded_filenames') != current_filenames or 
        st.session_state.selected_card != selected_card or 
        st.session_state.template_choice != template_choice or
        st.session_state.name_format != name_format):
        st.session_state.processing_done = False
        st.session_state.uploaded_filenames = current_filenames

if st.button("⚙️ تشغيل محرك الدمج والتنظيم والتنسيق المتقدم"):
    if uploaded_files:
        with st.spinner('جاري تجميع الملفات، تنظيف البيانات، الفرز الأبجدي، وإعداد التنسيق الشرطي...'):
            try:
                # تمرير قائمة الملفات المرفوعة إلى الدالة
                df_res = extract_and_clean_data(uploaded_files, selected_card, name_format)
                
                if not df_res.empty:
                    st.session_state.df_final = df_res
                    # إذا كان ملفاً واحداً نأخذ اسمه، وإذا كانت عدة ملفات نضع اسم الملف الأول مع كلمة (مدمج)
                    base_name = uploaded_files[0].name.rsplit('.', 1)[0]
                    if len(uploaded_files) > 1:
                        base_name += "_ومدمج"
                    
                    st.session_state.output_filename = base_name
                    st.session_state.selected_card = selected_card
                    st.session_state.template_choice = template_choice
                    st.session_state.name_format = name_format 
                    st.session_state.processing_done = True
                else:
                    st.error("لم يتم العثور على بيانات متوافقة في الملفات المرفوعة.")
            except Exception as e:
                st.error(f"خطأ غير متوقع أثناء المعالجة: {e}")
    else:
        st.warning("الرجاء رفع ملف واحد أو أكثر (docx أو xlsx) أولاً.")

if st.session_state.processing_done:
    df_final = st.session_state.df_final
    output_filename = st.session_state.output_filename
    used_card_type = st.session_state.selected_card
    used_template = st.session_state.template_choice
    used_name_format = st.session_state.name_format
    
    success_msg = "الاسم الثلاثي فقط" if used_name_format == "الاسم الثلاثي فقط (بدون لقب)" else "الاسم الكامل مع اللقب"
    st.success(f"✅ تمت عملية الدمج والتنظيم الأبجدي بنجاح لـ ({len(df_final)}) قيد (تم قصر الأسماء على: {success_msg}).")
    
    with st.spinner('جاري صياغة وهيكلة مستند Word النهائي...'):
        if used_template == "النموذج الأول (الأصلي المطور)":
            word_output = build_professional_word_report(df_final, output_filename, used_card_type)
        elif used_template == "النموذج الثاني (حجم 14 وحقلين فارغين)":
            word_output = build_professional_word_report_v2(df_final, output_filename, used_card_type)
        elif used_template == "النموذج الثالث (خط 16، عناوين 12، 4 أشهر)":
            word_output = build_professional_word_report_v3(df_final, output_filename, used_card_type)
        else:
            word_output = build_professional_word_report_v4(df_final, output_filename, used_card_type)
        
    st.download_button(
        label="📥 تحميل الكشف المدمج والمنسق (جاهز للطباعة فوراً)",
        data=word_output,
        file_name=f"كشف_جاهز_{output_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
