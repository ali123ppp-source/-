# import_streamlit_final.py
# نسخة نهائية ومتكاملة — دالة _parse_docx_to_list مدمجة ومحسّنة تماماً
# - معالجة vMerge عمودي/أفقي
# - دمج ذكي للأسطر المبعثرة (أرقام منفصلة، أسماء مقسومة)
# - تحليل HTML-like داخل الفقرات
# - سجل failed_lines و merge_groups و second-pass recovery
# - دعم ملفات docx, xlsx, csv للدمج
# - واجهة Streamlit كاملة لتحميل الملفات وتنزيل التقارير واللوجات
#
# تشغيل: streamlit run import_streamlit_final.py

import streamlit as st
import pandas as pd
from io import BytesIO, StringIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import csv
import traceback

# ---------------------------
# واجهة وإعدادات Streamlit
# ---------------------------
st.set_page_config(page_title="نظام كشوفات الوكلاء — نهائي", layout="wide")
st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { background-color: #1A5276; color: white; width: 100%; font-weight: bold; border-radius: 8px; font-size: 16px;}
    .report-box { background-color: #F4F6F7; padding: 12px; border-radius: 8px; border-right: 5px solid #1A5276; text-align: right; margin-bottom: 10px;}
    </style>
""", unsafe_allow_html=True)
st.markdown("<h2 style='text-align: right; color: #1A5276;'>نظام استخراج وتنسيق كشوفات الوكلاء — نهائي</h2>", unsafe_allow_html=True)

if "processing_done" not in st.session_state:
    st.session_state.processing_done = False
    st.session_state.df_final = None
    st.session_state.df_duplicates = None
    st.session_state.failed_lines = []
    st.session_state.merge_groups = []
    st.session_state.stats = {}
    st.session_state.output_filename = ""
    st.session_state.selected_card = ""

# ---------------------------
# أدوات تنسيق Word (ثابت)
# ---------------------------
def set_table_borders(table, color_hex="1A5276"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_vertical_text(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    text_dir = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="btLr"/>')
    tcPr.append(text_dir)

def set_cell_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = parse_xml(f'<w:noWrap {nsdecls("w")}/>') 
    tcPr.append(no_wrap)

def format_cell_advanced(cell, text, bold=False, color_rgb=None, size_pt=16, font_name="Calibri", align="center"):
    cell.text = str(text)
    p = cell.paragraphs[0]
    if align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left": p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    for run in p.runs:
        run.bold = bold
        if color_rgb: run.font.color.rgb = color_rgb
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        run.font.size = Pt(size_pt)

# ---------------------------
# دوال مساعدة
# ---------------------------
def classify_failure_reason(line):
    if not line or not line.strip():
        return "empty_line"
    low = line.lower()
    if '<td' in low or '<table' in low or '&lt;td' in low:
        return "html_table_unparsed"
    if not re.search(r'[\u0600-\u06FF]{2,}', line):
        return "no_arabic_name"
    if len(re.findall(r'\d', line)) < 1:
        return "no_numbers"
    return "ambiguous_format"

def row_cells_unique(row):
    seen = set()
    for cell in row.cells:
        tc = cell._tc
        if tc in seen:
            continue
        seen.add(tc)
        yield cell

def safe_int(val):
    try:
        return int(float(val))
    except:
        return 0

# ---------------------------
# الدالة الأساسية: _parse_docx_to_list (مُدمجة ومحسّنة)
# ---------------------------
def _parse_docx_to_list(file_obj, card_choice, params=None, stop_on_error=False, debug_limit=0):
    """
    Returns:
      parsed_records: list of dicts
      failed_lines: list of tuples (raw_line, reason, context_before, context_after, attempts)
      raw_count: int
      merge_groups: list of tuples (index_range, merged_text, reason)
      diagnostics: dict
    """
    if params is None:
        params = {}
    accept_card_min_len = params.get("accept_card_min_len", 4)
    accept_card_max_len = params.get("accept_card_max_len", 12)
    max_group_merge = params.get("max_group_merge", 5)
    enable_second_pass = params.get("enable_second_pass", True)

    doc = Document(file_obj)
    parsed_records = []
    failed_lines = []
    raw_lines = []
    merge_groups = []
    diagnostics = {"tables_read":0, "paras_read":0, "raw_lines_initial":0, "raw_lines_after_merge":0, "recovered_in_second_pass":0}

    # parse HTML-like table text inside paragraph
    def parse_html_table_text(text):
        tds = re.findall(r'<td[^>]*>(.*?)</td>', text, flags=re.DOTALL | re.IGNORECASE)
        tds = [re.sub(r'<.*?>', '', td).replace('\n',' ').strip() for td in tds]
        rows = []
        if not tds:
            return rows
        for cols in (8,7,6,5,4):
            if len(tds) % cols == 0:
                for i in range(0, len(tds), cols):
                    rows.append(' | '.join(tds[i:i+cols]))
                return rows
        rows.append(' | '.join(tds))
        return rows

    # 1) Read docx tables with robust vMerge handling
    for table in doc.tables:
        diagnostics["tables_read"] += 1
        try:
            max_cols = max(len(list(row_cells_unique(row))) for row in table.rows) if table.rows else 0
        except Exception:
            max_cols = 0
        n_rows = len(table.rows)
        if max_cols == 0 or n_rows == 0:
            continue

        grid = [["" for _ in range(max_cols)] for _ in range(n_rows)]
        vmerge_state = ["" for _ in range(max_cols)]

        for r_idx, row in enumerate(table.rows):
            cells = list(row_cells_unique(row))
            c_idx = 0
            for cell in cells:
                txt = cell.text.replace('\n', ' ').strip()
                vmerge_val = None
                try:
                    tcPr = cell._tc.tcPr
                    if tcPr is not None:
                        vm = tcPr.find(qn('w:vMerge'))
                        if vm is not None:
                            vmerge_val = vm.get(qn('w:val')) if vm.get(qn('w:val')) is not None else "continue"
                except Exception:
                    vmerge_val = None

                if vmerge_val and str(vmerge_val).lower() == 'continue':
                    grid[r_idx][c_idx] = vmerge_state[c_idx] if vmerge_state[c_idx] else txt
                else:
                    grid[r_idx][c_idx] = txt
                    if txt:
                        vmerge_state[c_idx] = txt
                c_idx += 1

            while c_idx < max_cols:
                grid[r_idx][c_idx] = vmerge_state[c_idx] if vmerge_state[c_idx] else ""
                c_idx += 1

        for r in range(n_rows):
            line = " | ".join([cell for cell in grid[r]])
            if line.strip():
                raw_lines.append(line)

    # 2) Read paragraphs and merge short fragments
    buffer_para = ""
    for para in doc.paragraphs:
        diagnostics["paras_read"] += 1
        txt = para.text.strip()
        if not txt:
            if buffer_para:
                raw_lines.append(buffer_para.strip())
                buffer_para = ""
            continue
        low = txt.lower()
        if '<table' in low or '<td' in low or '&lt;td' in low:
            rows = parse_html_table_text(txt)
            if rows:
                raw_lines.extend(rows)
            else:
                raw_lines.append(txt)
            continue
        if len(txt.split()) < 4:
            buffer_para += " " + txt
        else:
            if buffer_para:
                combined = (buffer_para + " " + txt).strip()
                raw_lines.append(combined)
                buffer_para = ""
            else:
                raw_lines.append(txt)
    if buffer_para:
        raw_lines.append(buffer_para.strip())

    diagnostics["raw_lines_initial"] = len(raw_lines)

    # 3) Pre-merge fragmented numeric runs and index-only lines
    arabic_re = re.compile(r'[\u0600-\u06FF]')
    numeric_only_re = re.compile(r'^\s*(?:\d+\s*)+$')
    short_nums_re = re.compile(r'^\s*(?:\d{1,3}\s+){1,}\d{1,3}\s*$')
    index_prefix_re = re.compile(r'^\s*\d+\s+')

    def merge_fragmented(lines):
        merged = []
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i].strip()
            if numeric_only_re.match(line) or (short_nums_re.match(line) and not arabic_re.search(line)):
                j = i
                nums = []
                while j < n and (numeric_only_re.match(lines[j].strip()) or (short_nums_re.match(lines[j].strip()) and not arabic_re.search(lines[j]))):
                    nums.append(lines[j].strip())
                    j += 1
                merged_nums = " ".join(nums)
                if merged and arabic_re.search(merged[-1]):
                    merged[-1] = merged[-1] + " " + merged_nums
                    merge_groups.append(((i-1, j-1), merged[-1], "numbers_to_prev"))
                elif j < n and arabic_re.search(lines[j]):
                    merged.append(lines[j].strip() + " " + merged_nums)
                    merge_groups.append(((i, j), merged[-1], "numbers_to_next"))
                    j += 1
                else:
                    merged.append(merged_nums)
                i = j
                continue

            if index_prefix_re.match(line) and not arabic_re.search(line) and i + 1 < n and arabic_re.search(lines[i+1]):
                merged.append(line + " " + lines[i+1].strip())
                merge_groups.append(((i, i+1), merged[-1], "index_to_next_name"))
                i += 2
                continue

            if arabic_re.search(line) and i + 1 < n and numeric_only_re.match(lines[i+1].strip()):
                merged.append(line + " " + lines[i+1].strip())
                merge_groups.append(((i, i+1), merged[-1], "name_then_numbers"))
                i += 2
                continue

            if arabic_re.search(line) and i + 1 < n and arabic_re.search(lines[i+1]) and len(line.split()) <= 2:
                if re.search(r'\d', lines[i+1]) or len(lines[i+1].split()) > 2:
                    merged.append(line + " " + lines[i+1].strip())
                    merge_groups.append(((i, i+1), merged[-1], "split_name_merge"))
                    i += 2
                    continue

            merged.append(line)
            i += 1

        final = []
        i = 0
        n = len(merged)
        while i < n:
            group = [merged[i]]
            j = i + 1
            while j < n and len(group) < max_group_merge:
                if not arabic_re.search(" ".join(group)) and arabic_re.search(merged[j]):
                    group.append(merged[j]); j += 1; break
                if numeric_only_re.match(merged[j]):
                    group.append(merged[j]); j += 1; continue
                break
            if len(group) > 1:
                merge_groups.append(((i, j-1), " ".join(group), "post_group_merge"))
            final.append(" ".join(group))
            i = j
        return final

    raw_lines = merge_fragmented(raw_lines)
    diagnostics["raw_lines_after_merge"] = len(raw_lines)

    # 4) Primary parse pass
    header_keywords = ["المركز", "اسم رب", "ملاحظات", "الوكيل", "الافراد", "الكلية", "المحجوبين", "FOOD", "نوع الوكالة", "ت رقم البطاقة", "اجمالي", "المجموع"]
    for idx, line in enumerate(raw_lines):
        attempts = 1
        try:
            line_clean = re.sub(r'[,\-،]', ' ', line)
            line_clean = re.sub(r'\s+', ' ', line_clean).strip()
            if any(re.search(r'^\s*' + re.escape(w) + r'\b', line_clean, flags=re.IGNORECASE) for w in header_keywords):
                if not (re.search(r'[\u0600-\u06FF]{2,}', line_clean) and re.search(r'\d', line_clean)):
                    continue

            all_nums = re.findall(r'\b0*\d{4,12}\b', line_clean)
            smalls_all = re.findall(r'\b\d{1,3}\b', line_clean)

            name_matches = re.findall(r'[\u0600-\u06FF]{2,}(?:\s+[\u0600-\u06FF]{2,})*', line_clean)
            name = None
            if name_matches:
                name = max(name_matches, key=len).strip()

            if not name:
                parts = [p.strip() for p in re.split(r'\||\t', line_clean) if p.strip()]
                for p in parts:
                    if re.search(r'[\u0600-\u06FF]{2,}', p):
                        name = p; break

            if not name:
                context_parts = []
                if idx > 0: context_parts.append(raw_lines[idx-1])
                context_parts.append(line_clean)
                if idx + 1 < len(raw_lines): context_parts.append(raw_lines[idx+1])
                combined = " ".join(context_parts)
                name_matches = re.findall(r'[\u0600-\u06FF]{2,}(?:\s+[\u0600-\u06FF]{2,})*', combined)
                if name_matches:
                    name = max(name_matches, key=len).strip()
                    merge_groups.append(((max(0, idx-1), min(len(raw_lines)-1, idx+1)), combined, "contextual_name_recover"))
                    attempts += 1

            if not name:
                reason = classify_failure_reason(line_clean)
                failed_lines.append((line, reason, raw_lines[idx-1] if idx>0 else "", raw_lines[idx+1] if idx+1<len(raw_lines) else "", attempts))
                continue

            cards = [n for n in all_nums if len(re.sub(r'^0+','',n)) >= accept_card_min_len or len(n) >= accept_card_min_len]
            old_card = cards[0] if cards else "غير متوفر"
            new_card = cards[-1] if len(cards) >= 2 else old_card
            selected_card = new_card if card_choice == "رقم البطاقة الحديث" else old_card

            idx_name = line_clean.find(name.split()[0]) if name else -1
            before = line_clean[:idx_name] if idx_name >= 0 else ""
            after = line_clean[idx_name + len(name):] if idx_name >= 0 else line_clean

            smalls_before = [int(n) for n in re.findall(r'\b\d{1,3}\b', before)]
            smalls_after = [int(n) for n in re.findall(r'\b\d{1,3}\b', after)]

            total = eligible = withheld = 0
            nums = smalls_after if len(smalls_after) >= 1 else smalls_before
            if len(nums) >= 3:
                total, eligible, withheld = nums[0], nums[1], nums[2]
            elif len(nums) == 2:
                total, eligible = nums[0], nums[1]; withheld = 0
            elif len(nums) == 1:
                total = nums[0]; eligible = withheld = 0
            else:
                all_smalls = [int(n) for n in re.findall(r'\b\d{1,3}\b', line_clean)]
                if len(all_smalls) >= 3:
                    total, eligible, withheld = all_smalls[-3], all_smalls[-2], all_smalls[-1]
                elif len(all_smalls) == 2:
                    total, eligible = max(all_smalls), min(all_smalls); withheld = 0
                else:
                    reason = classify_failure_reason(line_clean)
                    failed_lines.append((line, reason, raw_lines[idx-1] if idx>0 else "", raw_lines[idx+1] if idx+1<len(raw_lines) else "", attempts))
                    continue

            if re.search(r'^\s*(اجمالي|المجموع|مجموع)\b', name):
                continue

            parsed_records.append({
                "اسم رب الأسرة": name,
                "رقم البطاقة": selected_card,
                "الكلي": total,
                "محجوب": withheld,
                "مستحق": eligible
            })

            if debug_limit and len(parsed_records) >= debug_limit:
                break

        except Exception as e:
            failed_lines.append((line, f"exception:{str(e)}", raw_lines[idx-1] if idx>0 else "", raw_lines[idx+1] if idx+1<len(raw_lines) else "", attempts))
            if stop_on_error:
                raise

    # 5) Second-pass recovery
    if enable_second_pass and failed_lines:
        new_failed = []
        for entry in failed_lines:
            raw_line, reason, ctx_before, ctx_after, attempts = entry
            attempts += 1
            combined_candidates = []
            if ctx_before:
                combined_candidates.append(ctx_before + " " + raw_line)
            if ctx_after:
                combined_candidates.append(raw_line + " " + ctx_after)
            if ctx_before and ctx_after:
                combined_candidates.append(ctx_before + " " + raw_line + " " + ctx_after)
            recovered_flag = False
            for cand in combined_candidates:
                cand_clean = re.sub(r'[,\-،]', ' ', cand)
                name_matches = re.findall(r'[\u0600-\u06FF]{2,}(?:\s+[\u0600-\u06FF]{2,})*', cand_clean)
                if name_matches and re.search(r'\d', cand_clean):
                    name = max(name_matches, key=len).strip()
                    all_nums = re.findall(r'\b0*\d{4,12}\b', cand_clean)
                    cards = [n for n in all_nums if len(re.sub(r'^0+','',n)) >= accept_card_min_len or len(n) >= accept_card_min_len]
                    old_card = cards[0] if cards else "غير متوفر"
                    new_card = cards[-1] if len(cards) >= 2 else old_card
                    selected_card = new_card if card_choice == "رقم البطاقة الحديث" else old_card
                    smalls_all = [int(n) for n in re.findall(r'\b\d{1,3}\b', cand_clean)]
                    if len(smalls_all) >= 3:
                        total, eligible, withheld = smalls_all[-3], smalls_all[-2], smalls_all[-1]
                    elif len(smalls_all) == 2:
                        total, eligible = max(smalls_all), min(smalls_all); withheld = 0
                    elif len(smalls_all) == 1:
                        total = smalls_all[0]; eligible = withheld = 0
                    else:
                        continue
                    parsed_records.append({
                        "اسم رب الأسرة": name,
                        "رقم البطاقة": selected_card,
                        "الكلي": total,
                        "محجوب": withheld,
                        "مستحق": eligible
                    })
                    diagnostics["recovered_in_second_pass"] += 1
                    merge_groups.append(("second_pass", cand_clean, "second_pass_recovery"))
                    recovered_flag = True
                    break
            if not recovered_flag:
                new_failed.append((raw_line, reason, ctx_before, ctx_after, attempts))
        failed_lines = new_failed

    diagnostics["final_parsed"] = len(parsed_records)
    diagnostics["final_failed"] = len(failed_lines)
    diagnostics["merge_groups_count"] = len(merge_groups)

    return parsed_records, failed_lines, diagnostics["raw_lines_initial"], merge_groups, diagnostics

# ---------------------------
# قراءة Excel/CSV
# ---------------------------
def _parse_excel_to_list(file_obj, file_ext):
    records = []
    try:
        if file_ext == 'xlsx':
            df_in = pd.read_excel(file_obj, dtype=str)
        else:
            df_in = pd.read_csv(file_obj, dtype=str)
        df_in.columns = df_in.columns.astype(str).str.strip()
        col_name = next((c for c in df_in.columns if 'اسم' in c), df_in.columns[0])
        col_card = next((c for c in df_in.columns if 'بطاق' in c), df_in.columns[-1])
        col_total = next((c for c in df_in.columns if 'كلي' in c), None)
        col_eligible = next((c for c in df_in.columns if 'مستحق' in c), None)
        col_withheld = next((c for c in df_in.columns if 'محجوب' in c), None)

        for _, row in df_in.iterrows():
            name = str(row.get(col_name, "")).strip()
            if not name or name.lower() == 'nan' or 'اسم' in name: continue
            card = str(row.get(col_card, "")).strip()
            if not card or card.lower() == 'nan': card = "غير متوفر"
            records.append({
                "اسم رب الأسرة": name,
                "رقم البطاقة": card,
                "الكلي": safe_int(row.get(col_total, 0)) if col_total else 0,
                "محجوب": safe_int(row.get(col_withheld, 0)) if col_withheld else 0,
                "مستحق": safe_int(row.get(col_eligible, 0)) if col_eligible else 0
            })
    except Exception as e:
        st.warning(f"تنبيه: حدث خطأ أثناء تحليل ملف الإكسل/الجدول الإضافي ({e})")
    return records

# ---------------------------
# استخراج ودمج وتنظيف (ماستر)
# ---------------------------
def extract_and_clean_data(file_obj, card_choice, file_obj2=None, file2_ext=None):
    parsed_primary, failed_primary, raw_count_primary, merge_groups_primary, diag_primary = _parse_docx_to_list(file_obj, card_choice)
    parsed_secondary = []
    failed_secondary = []
    raw_count_secondary = 0
    merge_groups_secondary = []
    diag_secondary = {}

    if file_obj2:
        if file2_ext == 'docx':
            parsed_secondary, failed_secondary, raw_count_secondary, merge_groups_secondary, diag_secondary = _parse_docx_to_list(file_obj2, card_choice)
        elif file2_ext in ['xlsx', 'csv']:
            parsed_secondary = _parse_excel_to_list(file_obj2, file2_ext)
            raw_count_secondary = len(parsed_secondary)
        else:
            pass

    all_parsed = parsed_primary + parsed_secondary
    all_failed = failed_primary + failed_secondary
    all_merge_groups = merge_groups_primary + merge_groups_secondary

    df = pd.DataFrame(all_parsed)
    df_duplicates = pd.DataFrame()
    if not df.empty:
        duplicate_mask = df.duplicated(keep='first')
        df_duplicates = df[duplicate_mask].copy()
        if not df_duplicates.empty:
            df_duplicates = df_duplicates.sort_values(by="اسم رب الأسرة").reset_index(drop=True)
            df_duplicates.insert(0, "ت", df_duplicates.index + 1)
        df = df.drop_duplicates(keep='first')
        df = df.sort_values(by="اسم رب الأسرة").reset_index(drop=True)
        df.insert(0, "ت", df.index + 1)

    stats = {
        "raw_lines_primary": raw_count_primary,
        "raw_lines_secondary": raw_count_secondary,
        "parsed_count": len(all_parsed),
        "failed_count": len(all_failed),
        "duplicates_count": len(df_duplicates),
        "merge_groups_count": len(all_merge_groups),
        "diag_primary": diag_primary,
        "diag_secondary": diag_secondary
    }
    return df, df_duplicates, all_failed, all_merge_groups, stats

# ---------------------------
# بناء تقرير Word (ثابت)
# ---------------------------
def build_professional_word_report(df, filename_base, card_choice, is_duplicate_report=False):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.3)
        section.right_margin = Cm(0.3)

    clean_name = filename_base
    for w in ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]: clean_name = clean_name.replace(w, "")
    clean_name = re.sub(r'[a-zA-Z]', '', clean_name)
    clean_name = re.sub(r'[\-_+_.]', '', clean_name)
    clean_name = " ".join(clean_name.split())

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    report_title = f"كشف المتكررات المعزولة: {clean_name}" if is_duplicate_report else f"الكشف الإحصائي المنسق للوكيل: {clean_name}"
    title_run = title_p.add_run(report_title)
    title_run.font.name = "Segoe UI Semibold"
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(220, 20, 60) if is_duplicate_report else RGBColor(26, 82, 118)

    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("صفحة ")
    f_run.font.size = Pt(10)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    f_run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

    headers = ["ت", "اسم رب الأسرة", "حقل فارغ", "الكلي", "مستحق", "محجوب", card_choice, "ملاحظات"]
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    border_color = "DC143C" if is_duplicate_report else "1A5276"
    set_table_borders(table, color_hex=border_color)
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    if not df.empty:
        max_name_len = max(df["اسم رب الأسرة"].astype(str).str.len().max(), 15)
    else:
        max_name_len = 15
    dynamic_name_width = Cm(max_name_len * 0.22 + 0.5)
    col_widths = [Cm(0.9), dynamic_name_width, Cm(0.44), Cm(0.9), Cm(0.9), Cm(0.9), Cm(3.0), Cm(2.19)]
    THEME_COLOR = RGBColor(220, 20, 60) if is_duplicate_report else RGBColor(26, 82, 118)

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        if i in [3, 4, 5]:
            set_cell_vertical_text(hdr_cells[i])
            format_cell_advanced(hdr_cells[i], title, bold=True, size_pt=12, font_name="Segoe UI Semibold", align="center", color_rgb=THEME_COLOR)
        else:
            cell_align = "left" if i == 1 else "center"
            format_cell_advanced(hdr_cells[i], title, bold=True, size_pt=14, font_name="Segoe UI Semibold", align=cell_align, color_rgb=THEME_COLOR)

    HEX_ELEGANT_BLUE = "D4E6F1"
    HEX_LIGHT_BLUE = "EBF5FB"
    HEX_LIGHT_RED = "FADBD8"
    HEX_LIGHT_GREEN = "E8F8F5"
    HEX_ALERT_RED = "EC7063"

    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        r_trPr = table.rows[idx+1]._tr.get_or_add_trPr()
        r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        is_eligible_zero = int(row["مستحق"]) == 0
        for i in range(8): row_cells[i].width = col_widths[i]
        set_cell_no_wrap(row_cells[1])
        for i in range(8):
            val, text_color, cell_align, font_size = "", None, "center", 16
            if i == 0: val = row["ت"]
            elif i == 1:
                val = row["اسم رب الأسرة"]
                cell_align = "left"
            elif i == 2: val = "x" if is_eligible_zero else ""
            elif i == 3: val = row["الكلي"]
            elif i == 4: val = row["مستحق"]
            elif i == 5:
                val = row["محجوب"]
                font_size = 14
            elif i == 6: val = row["رقم البطاقة"]
            elif i == 7:
                val = "محجوب" if is_eligible_zero else ""
                if is_eligible_zero:
                    text_color = RGBColor(203, 67, 53)
                    font_size = 12
            format_cell_advanced(row_cells[i], val, size_pt=font_size, font_name="Calibri", color_rgb=text_color, align=cell_align)
            if is_eligible_zero: set_cell_background(row_cells[i], HEX_ALERT_RED)
            else:
                if i == 0: set_cell_background(row_cells[i], HEX_ELEGANT_BLUE)
                elif i == 3: set_cell_background(row_cells[i], HEX_LIGHT_BLUE)
                elif i == 4: set_cell_background(row_cells[i], HEX_LIGHT_GREEN)
                elif i == 5: set_cell_background(row_cells[i], HEX_LIGHT_RED)

    total_all = df["الكلي"].astype(int).sum() if not df.empty else 0
    total_eligible = df["مستحق"].astype(int).sum() if not df.empty else 0
    total_withheld = df["محجوب"].astype(int).sum() if not df.empty else 0

    doc.add_paragraph()
    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stats_p.paragraph_format.element.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    stats_text = (f"العدد الكلي للافراد = {total_all}\n"
                  f"العدد الكلي للمستحقين = {total_eligible}\n"
                  f"العدد الكلي للمحجوبين = {total_withheld}")
    stats_run = stats_p.add_run(stats_text)
    stats_run.font.name = "Segoe UI Semibold"
    stats_run.font.size = Pt(13)
    stats_run.bold = True
    stats_run.font.color.rgb = THEME_COLOR

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------
# واجهة المستخدم: رفع ملفات وتشغيل المحرك
# ---------------------------
col1, col2 = st.columns(2)
with col2:
    st.markdown("<h4 style='text-align: right;'>📂 رفع الكشف الأساسي (Word)</h4>", unsafe_allow_html=True)
    uploaded_file = st.file_uploader("ارفع كشف الوكلاء الأساسي (docx)", type=['docx'], key="doc_input_master", label_visibility="collapsed")
with col1:
    st.markdown("<h4 style='text-align: right;'>➕ ملف للدمج (اختياري)</h4>", unsafe_allow_html=True)
    uploaded_file2 = st.file_uploader("ارفع جدول (docx/xlsx/csv) لدمجه أبجدياً (اختياري)", type=['docx', 'xlsx', 'csv'], key="doc_input_merge", label_visibility="collapsed")

st.markdown("<br>", unsafe_allow_html=True)
selected_card = st.radio("اختر نوع رقم البطاقة المراد اعتماده في الكشف:", ["رقم البطاقة القديم", "رقم البطاقة الحديث"], index=0, horizontal=True)
st.markdown("<br>", unsafe_allow_html=True)

if uploaded_file:
    current_filename = uploaded_file.name.rsplit('.', 1)[0]
    if st.session_state.output_filename != current_filename or st.session_state.selected_card != selected_card:
        st.session_state.processing_done = False

if st.button("🚀 تشغيل المحرك (استخراج ودمج وفرز)"):
    if not uploaded_file:
        st.warning("الرجاء رفع ملف الكشف الأساسي (Word) أولاً.")
    else:
        with st.spinner('يتم الآن معالجة الملف...'):
            try:
                file2_ext = uploaded_file2.name.split('.')[-1].lower() if uploaded_file2 else None
                df_res, df_dup, failed_lines, merge_groups, stats = extract_and_clean_data(uploaded_file, selected_card, uploaded_file2, file2_ext)
                st.session_state.df_final = df_res
                st.session_state.df_duplicates = df_dup
                st.session_state.failed_lines = failed_lines
                st.session_state.merge_groups = merge_groups
                st.session_state.stats = stats
                st.session_state.output_filename = uploaded_file.name.rsplit('.', 1)[0]
                st.session_state.selected_card = selected_card
                st.session_state.processing_done = True

                if df_res.empty:
                    st.error("لم يتم العثور على سجلات قابلة للاستخراج. راجع الملف أو اطلع على لوج الفشل.")
                else:
                    st.success(f"استخراج أولي: تم استخراج {len(df_res)} سجلّاً؛ فشل في استخراج {len(failed_lines)} سطر؛ مجموعات دمج مُسجلة {len(merge_groups)}.")
            except Exception as e:
                st.error(f"خطأ غير متوقع أثناء المعالجة: {e}")
                st.error(traceback.format_exc())

# ---------------------------
# بعد المعالجة: تنزيل التقارير وملفات اللوج
# ---------------------------
if st.session_state.processing_done:
    df_final = st.session_state.df_final
    df_duplicates = st.session_state.df_duplicates
    failed_lines = st.session_state.failed_lines
    merge_groups = st.session_state.merge_groups
    stats = st.session_state.stats
    output_filename = st.session_state.output_filename
    used_card_type = st.session_state.selected_card

    st.balloons()
    merge_text = " (شاملة البيانات المدمجة)" if uploaded_file2 else ""
    st.success(f"🏁 انتهت المعالجة — تم استخراج ({len(df_final)}) قيد صافٍ أبجدياً{merge_text}.")

    with st.spinner('جاري بناء ملف Word...'):
        word_output = build_professional_word_report(df_final, output_filename, used_card_type, is_duplicate_report=False)

    st.download_button(
        label=f"📥 تحميل الكشف الأساسي المدمج ({len(df_final)} اسم)",
        data=word_output,
        file_name=f"كشف_منسق_مدمج_{output_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    if not df_duplicates.empty:
        st.warning(f"⚠️ تم عزل ({len(df_duplicates)}) قيد كـ تكرارات.")
        with st.spinner('جاري بناء ملف التكرارات...'):
            word_output_dup = build_professional_word_report(df_duplicates, output_filename, used_card_type, is_duplicate_report=True)
        st.download_button(
            label="🚨 تحميل ملف القيود المتكررة المعزولة",
            data=word_output_dup,
            file_name=f"سجلات_متكررة_معزولة_{output_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.markdown("### إحصاءات الاستخراج")
    st.write({
        "عدد الأسطر الخام (أساسي)": stats.get("raw_lines_primary", 0),
        "عدد الأسطر الخام (ثانوي)": stats.get("raw_lines_secondary", 0),
        "عدد السجلات المستخرجة": stats.get("parsed_count", 0),
        "عدد الأسطر الفاشلة": stats.get("failed_count", 0),
        "عدد التكرارات المعزولة": stats.get("duplicates_count", 0),
        "عدد مجموعات الدمج المسجلة": stats.get("merge_groups_count", 0),
    })

    if failed_lines:
        st.markdown("### أمثلة على الأسطر التي فشل استخراجها (مع سبب الفشل)")
        sample = failed_lines[:200]
        df_failed = pd.DataFrame(sample, columns=["السطر الخام", "سبب الفشل", "السطر السابق", "السطر التالي", "محاولات"])
        st.dataframe(df_failed)

        csv_buffer = StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(["السطر الخام", "سبب الفشل", "السطر السابق", "السطر التالي", "محاولات"])
        for ln, reason, before, after, attempts in failed_lines:
            writer.writerow([ln, reason, before, after, attempts])
        csv_buffer.seek(0)
        st.download_button(
            label="📥 تحميل لوج الأسطر الفاشلة (CSV)",
            data=csv_buffer.getvalue(),
            file_name=f"failed_lines_{output_filename}.csv",
            mime="text/csv"
        )

    if merge_groups:
        st.markdown("### أمثلة على مجموعات الدمج (merge groups)")
        mg_sample = merge_groups[:200]
        df_mg = pd.DataFrame([(str(idx_pair), merged_text, reason) for idx_pair, merged_text, reason in mg_sample],
                             columns=["فهرس/نطاق", "النص المدموج", "السبب"])
        st.dataframe(df_mg)

        csv_buf2 = StringIO()
        writer2 = csv.writer(csv_buf2)
        writer2.writerow(["فهرس/نطاق", "النص المدموج", "السبب"])
        for idx_pair, merged_text, reason in merge_groups:
            writer2.writerow([str(idx_pair), merged_text, reason])
        csv_buf2.seek(0)
        st.download_button(
            label="📥 تحميل لوج مجموعات الدمج (CSV)",
            data=csv_buf2.getvalue(),
            file_name=f"merge_groups_{output_filename}.csv",
            mime="text/csv"
        )

    st.markdown("### توصيات لتحسين التغطية إلى 100%")
    st.markdown("""
    - افتح المستند في Word واحفظه كملف جديد (حفظ كـ docx) للتخلص من HTML داخل الفقرات.
    - تأكد أن أسماء الأسر ليست مقسومة على خلايا متعددة داخل الجدول؛ إن وُجدت، أعد ترتيب الجدول أو زودني بلوج الفشل.
    - حمّل أول 50 سطرًا من failed_lines و أول 50 مجموعة من merge_groups وسأزوّدك بقواعد استثنائية دقيقة.
    """)

# نهاية الملف
